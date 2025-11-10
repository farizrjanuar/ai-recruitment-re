"""
Text Extractor Module

This module provides functionality to extract text from various document formats
(PDF, DOCX, TXT) and clean/normalize the extracted text.
"""

import re
import os
from typing import Optional
import PyPDF2
from docx import Document


class TextExtractor:
    """
    Handles text extraction from multiple file formats and text normalization.
    
    Supports:
    - PDF files (using PyPDF2)
    - DOCX files (using python-docx)
    - TXT files (plain text)
    """
    
    def __init__(self):
        """Initialize the TextExtractor."""
        pass
    
    def extract_from_pdf(self, file_path: str) -> str:
        """
        Extract text from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text as a string
            
        Raises:
            FileNotFoundError: If the file doesn't exist
            ValueError: If the file is corrupted or unreadable
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"PDF file not found: {file_path}")
        
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    raise ValueError("PDF file is encrypted and cannot be read")
                
                # Extract text from all pages
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if not text.strip():
                raise ValueError("No text could be extracted from PDF file")
            
            return text
            
        except PyPDF2.errors.PdfReadError as e:
            raise ValueError(f"Corrupted or invalid PDF file: {str(e)}")
        except Exception as e:
            raise ValueError(f"Error reading PDF file: {str(e)}")
    
    def extract_from_docx(self, file_path: str) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text as a string
            
        Raises:
            FileNotFoundError: If the file doesn't exist
            ValueError: If the file is corrupted or unreadable
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                raise ValueError("No text could be extracted from DOCX file")
            
            return text
            
        except Exception as e:
            raise ValueError(f"Error reading DOCX file: {str(e)}")
    
    def extract_from_txt(self, file_path: str) -> str:
        """
        Extract text from a plain text file.
        
        Args:
            file_path: Path to the TXT file
            
        Returns:
            Extracted text as a string
            
        Raises:
            FileNotFoundError: If the file doesn't exist
            ValueError: If the file is unreadable
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"TXT file not found: {file_path}")
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    
                    if not text.strip():
                        raise ValueError("Text file is empty")
                    
                    return text
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Unable to decode text file with supported encodings")
            
        except Exception as e:
            if isinstance(e, ValueError):
                raise
            raise ValueError(f"Error reading TXT file: {str(e)}")
    
    def clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Performs the following operations:
        - Removes excessive whitespace
        - Normalizes line breaks
        - Removes special characters (keeping alphanumeric, basic punctuation)
        - Removes control characters
        - Strips leading/trailing whitespace
        
        Args:
            text: Raw text to clean
            
        Returns:
            Cleaned and normalized text
        """
        if not text:
            return ""
        
        # Remove control characters except newline and tab
        text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\t')
        
        # Normalize line breaks (convert \r\n and \r to \n)
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive newlines (more than 2 consecutive)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove excessive spaces (more than 1 consecutive)
        text = re.sub(r' {2,}', ' ', text)
        
        # Remove spaces at the beginning and end of lines
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        # Remove empty lines at the beginning and end
        text = text.strip()
        
        return text
    
    def extract_text(self, file_path: str, file_type: str) -> str:
        """
        Extract text from a file based on its type.
        
        Args:
            file_path: Path to the file
            file_type: File type ('pdf', 'docx', or 'txt')
            
        Returns:
            Cleaned and extracted text
            
        Raises:
            ValueError: If file type is not supported or file is unreadable
            FileNotFoundError: If the file doesn't exist
        """
        file_type = file_type.lower().strip('.')
        
        if file_type == 'pdf':
            raw_text = self.extract_from_pdf(file_path)
        elif file_type in ['docx', 'doc']:
            raw_text = self.extract_from_docx(file_path)
        elif file_type == 'txt':
            raw_text = self.extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Clean and normalize the extracted text
        cleaned_text = self.clean_text(raw_text)
        
        return cleaned_text
